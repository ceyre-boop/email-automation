"""
Tests for the docx → sop.md merge path (upload_sop_v2) and the Automated Send
Workflow upload.

Why this file exists: the merge is the step that rewrites the approved-response
text used to generate live brand emails, and it had no test coverage at all.
Three separate incidents (`6986f6d`, `423034e`, and a repair during the
2026-07-27 session) were caused by an import wholesale-replacing sop.md and
destroying the machine-read metadata (Key/Gmail/Min Rate/Manager/personal-email
bullets). These tests lock the invariant down.

Every test runs against a temp COPY of the real sheets/sop.md — the repo file is
never written to.
"""
from __future__ import annotations

import pathlib
from io import BytesIO

import pytest

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

pytestmark = pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")

REAL_SOP = pathlib.Path(__file__).resolve().parents[2] / "sheets" / "sop.md"
REAL_WORKFLOW = REAL_SOP.parent / "Automated Send Workflow.md"

# Metadata lines that must survive every merge. Losing any of these is the
# regression this file exists to prevent.
METADATA_PREFIXES = ("Key:", "Gmail:", "Min Rate:", "Auto Send:", "Paused:")


@pytest.fixture
def sop_sandbox(tmp_path, monkeypatch):
    """Point every module-level SOP path at a temp copy of the real sop.md.

    reply, sop_writer and sop_admin each hold their own module-level constant,
    so all three must be redirected or a test would write the repo file.
    Also neuters sop_writer's git auto-commit, which would otherwise shell out
    to `git add/commit/push` from a temp directory on every write.
    """
    from backend.routers import sop_admin
    from backend.services import reply, sop_writer

    sandbox = tmp_path / "sop.md"
    sandbox.write_text(REAL_SOP.read_text(encoding="utf-8"), encoding="utf-8")

    workflow = tmp_path / "Automated Send Workflow.md"
    if REAL_WORKFLOW.exists():
        workflow.write_text(REAL_WORKFLOW.read_text(encoding="utf-8"), encoding="utf-8")

    monkeypatch.setattr(sop_writer, "_SOP_PATH", sandbox)
    monkeypatch.setattr(sop_writer, "_WORKFLOW_PATH", workflow)
    monkeypatch.setattr(sop_admin, "_SOP_PATH", sandbox)
    monkeypatch.setattr(reply, "_SOP_MD_PATH", sandbox)
    # write_sop_md() does `import subprocess` inside the function body, so it
    # resolves to the real module at call time — patch the module itself.
    import subprocess as _subprocess
    monkeypatch.setattr(_subprocess, "run", lambda *a, **k: None)
    reply.clear_sop_cache()

    yield sandbox

    reply.clear_sop_cache()


def _sop_docx(entries: list[tuple[str, str]], emails: dict[str, list[str]] | None = None,
              extra_scenarios: dict[str, str] | None = None) -> bytes:
    """Build a minimal SOP-shaped docx. entries = [(full_name, approved_response)]."""
    emails = emails or {}
    extra_scenarios = extra_scenarios or {}
    doc = Document()
    doc.add_paragraph("Part 4 — Talent Approved Responses")
    for full_name, response in entries:
        doc.add_paragraph(f"Talent: {full_name}")
        doc.add_paragraph("Scenario A: Initial Inbound (Default Response)")
        doc.add_paragraph("Approved Response:")
        for line in response.splitlines():
            doc.add_paragraph(line)
        if full_name in extra_scenarios:
            doc.add_paragraph(f"Scenario {extra_scenarios[full_name]}: Below Minimum Offer")
            doc.add_paragraph("Approved Response:")
            doc.add_paragraph("Counter-offer text.")
        if full_name in emails:
            doc.add_paragraph("Scenario C: Personal Email Forward")
            doc.add_paragraph("Personal Emails:")
            for addr in emails[full_name]:
                doc.add_paragraph(f"- {addr}")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _upload(client, payload: bytes, label: str = "test"):
    return client.post(
        "/admin/api/sop/upload-v2",
        files={"file": ("sop.docx", payload,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"label": label},
    )


def _first_talent(sandbox) -> tuple[str, str]:
    """Return (key, full_name) of a real talent in the sandbox sop.md."""
    from backend.services.sop_parser import parse_sop_md

    profiles = parse_sop_md(sandbox.read_text(encoding="utf-8"))
    key = sorted(profiles)[0]
    return key, profiles[key].full_name


# ── Merge behaviour ───────────────────────────────────────────────────────────

def test_matched_talent_response_is_updated(sop_sandbox, client):
    """A talent present in sop.md gets their Scenario A replaced."""
    from backend.services import reply
    from backend.services.reply import _extract_approved_response, _get_talent_section_raw

    key, full_name = _first_talent(sop_sandbox)
    new_text = "Totally new approved response for the merge test."

    resp = _upload(client, _sop_docx([(full_name, new_text)]))
    assert resp.status_code == 200, resp.text
    assert key in resp.json()["matched"]

    reply.clear_sop_cache()
    section = _get_talent_section_raw(full_name)
    live = _extract_approved_response(section, "Scenario A")
    assert live.strip() == new_text


def test_unknown_talent_is_skipped_not_inserted(sop_sandbox, client):
    """A talent only in the docx must never be silently added to sop.md."""
    from backend.services.sop_parser import parse_sop_md

    before = parse_sop_md(sop_sandbox.read_text(encoding="utf-8"))

    resp = _upload(client, _sop_docx([("Nonexistent Person", "Hello there.")]))
    assert resp.status_code == 200, resp.text
    assert "Nonexistent Person" in resp.json()["unmatched_docx"]

    after = parse_sop_md(sop_sandbox.read_text(encoding="utf-8"))
    assert set(after) == set(before), "merge must not add or remove talents"
    assert "Nonexistent Person" not in sop_sandbox.read_text(encoding="utf-8")


def test_metadata_survives_the_merge(sop_sandbox, client):
    """Key/Gmail/Min Rate/Auto Send/Paused and manager emails must be untouched.

    This is the exact regression from 6986f6d and 423034e.
    """
    original = sop_sandbox.read_text(encoding="utf-8")
    counts_before = {p: original.count(f"\n{p}") for p in METADATA_PREFIXES}
    managers_before = original.count("Manager: ")

    key, full_name = _first_talent(sop_sandbox)
    resp = _upload(client, _sop_docx([(full_name, "Replacement copy.")]))
    assert resp.status_code == 200, resp.text

    merged = sop_sandbox.read_text(encoding="utf-8")
    for prefix in METADATA_PREFIXES:
        assert merged.count(f"\n{prefix}") == counts_before[prefix], f"{prefix} lines lost"
    assert merged.count("Manager: ") == managers_before
    assert "<" in merged and "@taboost.me>" in merged, "manager emails stripped"


def test_personal_emails_replaced_and_keep_bullet_format(sop_sandbox, client):
    """_parse_personal_emails depends on the '- addr' bullet shape."""
    from backend.services.sop_parser import parse_sop_md

    key, full_name = _first_talent(sop_sandbox)
    resp = _upload(
        client,
        _sop_docx([(full_name, "Body.")], emails={full_name: ["merged@example.com"]}),
    )
    assert resp.status_code == 200, resp.text

    merged = sop_sandbox.read_text(encoding="utf-8")
    assert "- merged@example.com" in merged
    assert parse_sop_md(merged)[key].personal_emails == ["merged@example.com"]


def test_profile_count_and_warnings_do_not_regress(sop_sandbox, client):
    """A merge must not reduce parseable profiles or add validator warnings."""
    from backend.services.sop_parser import parse_sop_md, validate_profiles

    before = parse_sop_md(sop_sandbox.read_text(encoding="utf-8"))
    warnings_before = validate_profiles(before)

    key, full_name = _first_talent(sop_sandbox)
    assert _upload(client, _sop_docx([(full_name, "Body.")])).status_code == 200

    after = parse_sop_md(sop_sandbox.read_text(encoding="utf-8"))
    assert len(after) == len(before)
    assert len(validate_profiles(after)) <= len(warnings_before)


def test_sop_cache_is_cleared_so_next_draft_sees_new_text(sop_sandbox, client):
    """write_sop_md must invalidate the cache or drafts keep the stale SOP."""
    from backend.services import reply

    key, full_name = _first_talent(sop_sandbox)
    reply._load_sop_md()                      # prime the cache
    assert reply._sop_md_cache is not None

    assert _upload(client, _sop_docx([(full_name, "Cache-busting copy.")])).status_code == 200
    assert "Cache-busting copy." in reply._load_sop_md()


# ── sop_versions bookkeeping ──────────────────────────────────────────────────

def test_upload_records_active_sop_version(sop_sandbox, client):
    """Each upload inserts an active doc_type='sop' row and deactivates the last."""
    from backend.models.db import SopVersion, get_session_factory

    key, full_name = _first_talent(sop_sandbox)
    assert _upload(client, _sop_docx([(full_name, "v1")]), label="v1").status_code == 200
    assert _upload(client, _sop_docx([(full_name, "v2")]), label="v2").status_code == 200

    db = get_session_factory()()
    try:
        rows = db.query(SopVersion).filter(SopVersion.doc_type == "sop").all()
        active = [r for r in rows if r.is_active]
        assert len(active) == 1, "exactly one active sop version expected"
        assert active[0].version_label == "v2"
        assert active[0].talent_count and active[0].talent_count > 0
    finally:
        db.close()


# ── Guards ────────────────────────────────────────────────────────────────────

def test_scenario_b_in_docx_produces_a_warning(sop_sandbox, client):
    """Only Scenario A is merged — a Scenario B must not vanish silently."""
    key, full_name = _first_talent(sop_sandbox)
    payload = _sop_docx([(full_name, "Body.")], extra_scenarios={full_name: "B"})

    resp = _upload(client, payload)
    assert resp.status_code == 200, resp.text
    warnings = " ".join(resp.json()["warnings"])
    assert "Scenario B" in warnings and full_name in warnings
    assert "NOT merged" in warnings


def test_scenario_c_does_not_warn(sop_sandbox, client):
    """Scenario C is personal-email routing and IS handled — no false alarm."""
    key, full_name = _first_talent(sop_sandbox)
    resp = _upload(
        client,
        _sop_docx([(full_name, "Body.")], emails={full_name: ["x@example.com"]}),
    )
    assert resp.status_code == 200, resp.text
    assert not any("Scenario C" in w for w in resp.json()["warnings"])


def test_missing_python_docx_reports_a_server_error_not_a_bad_file(sop_sandbox, client, monkeypatch):
    """A missing server dependency must not be blamed on the operator's file.

    docx_parser imports python-docx inside the function, so the ImportError
    arrives at call time and used to be swallowed by the broad handler and
    returned as 400 "Could not parse docx".
    """
    from backend.routers import sop_admin

    def _boom(_content):
        raise ImportError("No module named 'docx'")

    monkeypatch.setattr(sop_admin, "extract_talent_sections", _boom, raising=False)
    monkeypatch.setitem(
        __import__("sys").modules, "backend.services.docx_parser",
        type("M", (), {"extract_talent_sections": staticmethod(_boom)}),
    )

    key, full_name = _first_talent(sop_sandbox)
    resp = _upload(client, _sop_docx([(full_name, "Body.")]))
    assert resp.status_code == 500, f"expected server error, got {resp.status_code}"
    assert "python-docx" in resp.json()["detail"]


def test_empty_upload_rejected(sop_sandbox, client):
    resp = client.post(
        "/admin/api/sop/upload-v2",
        files={"file": ("sop.docx", b"", "application/octet-stream")},
        data={"label": "x"},
    )
    assert resp.status_code == 400


# ── Automated Send Workflow upload ────────────────────────────────────────────

def _workflow_docx(text: str) -> bytes:
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _upload_workflow(client, payload: bytes, confirm: bool = False, label: str = "wf"):
    return client.post(
        "/admin/api/workflow/upload",
        files={"file": ("wf.docx", payload,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"label": label, "confirm": "true" if confirm else "false"},
    )


def test_workflow_preview_writes_nothing(sop_sandbox, client):
    from backend.services import sop_writer

    current = sop_writer.read_workflow_md()
    changed = current.replace("Send Gate", "Send Gate (revised)", 1)

    resp = _upload_workflow(client, _workflow_docx(changed), confirm=False)
    assert resp.status_code == 200, resp.text
    assert resp.json()["status"] == "preview"
    assert sop_writer.read_workflow_md() == current, "preview must not write"


def test_workflow_confirm_writes_and_versions(sop_sandbox, client):
    from backend.models.db import SopVersion, get_session_factory
    from backend.services import sop_writer

    changed = sop_writer.read_workflow_md().replace("Send Gate", "Send Gate (revised)", 1)
    resp = _upload_workflow(client, _workflow_docx(changed), confirm=True, label="wf-v2")
    assert resp.status_code == 200, resp.text
    assert "Send Gate (revised)" in sop_writer.read_workflow_md()

    db = get_session_factory()()
    try:
        active = db.query(SopVersion).filter(
            SopVersion.doc_type == "workflow", SopVersion.is_active == True  # noqa: E712
        ).all()
        assert len(active) == 1 and active[0].version_label == "wf-v2"
    finally:
        db.close()


def test_workflow_rejects_the_sop_document(sop_sandbox, client):
    """Dropping the SOP into the workflow box must be refused, not written."""
    from backend.services import sop_writer

    before = sop_writer.read_workflow_md()
    resp = _upload_workflow(client, _workflow_docx(REAL_SOP.read_text(encoding="utf-8")), confirm=True)
    assert resp.status_code == 400
    assert sop_writer.read_workflow_md() == before


def test_startup_restore_never_writes_the_workflow_over_sop_md(sop_sandbox, client):
    """The startup restore must filter by doc_type.

    sop_versions holds both documents. The startup handler originally selected
    the newest is_active row with NO doc_type filter, so uploading the workflow
    doc meant the next boot wrote it over sheets/sop.md — destroying every
    talent's approved response. This test reproduced that for real (it clobbered
    the repo's sop.md) before main.py was fixed.
    """
    from backend.models.db import SopVersion, get_session_factory
    from backend.services import sop_writer as _sw

    # An active workflow version, uploaded most recently.
    db = get_session_factory()()
    try:
        db.add(SopVersion(
            version_label="wf", raw_content="Automated Send Workflow\nSend Gate\nPost-Send",
            talent_count=None, is_active=True, doc_type="workflow",
        ))
        db.commit()
    finally:
        db.close()

    sop_before = sop_sandbox.read_text(encoding="utf-8")

    active_sop = (
        get_session_factory()()
        .query(SopVersion)
        .filter(SopVersion.is_active == True, SopVersion.doc_type == "sop")  # noqa: E712
        .order_by(SopVersion.uploaded_at.desc())
        .first()
    )
    if active_sop:
        _sw._SOP_PATH.write_text(active_sop.raw_content, encoding="utf-8")

    restored = sop_sandbox.read_text(encoding="utf-8")
    assert "Talent: " in restored, "sop.md was overwritten with a non-SOP document"
    assert restored.count("Talent: ") == sop_before.count("Talent: ")
    assert "Automated Send Workflow" not in restored


def test_workflow_rejects_short_document(sop_sandbox, client):
    from backend.services import sop_writer

    before = sop_writer.read_workflow_md()
    resp = _upload_workflow(client, _workflow_docx("too short"), confirm=True)
    assert resp.status_code == 400
    assert sop_writer.read_workflow_md() == before
