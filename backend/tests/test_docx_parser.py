"""
Tests for backend/services/docx_parser.py — the rich-text DOCX → markdown extractor.

These tests create minimal .docx files in memory using python-docx so they don't
depend on any external fixture files.  python-docx must be installed on the test
runner; the test is skipped with a clear message if it's not.
"""
from __future__ import annotations

from io import BytesIO

import pytest

try:
    from docx import Document
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

pytestmark = pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")


def _make_docx(**sections) -> bytes:
    """
    Build a minimal SOP-style .docx in memory.

    Keyword args control what paragraphs are written:
      talent_name  – str, written as "Talent: <name>"
      approved     – str, written after "Approved Response:"
      emails       – list[str], written after "Personal Emails:"
    """
    doc = Document()
    if sections.get("talent_name"):
        doc.add_paragraph(f"Talent: {sections['talent_name']}")
    if sections.get("approved") is not None:
        doc.add_paragraph("Scenario A: Initial Inbound")
        doc.add_paragraph("Approved Response:")
        for line in sections["approved"].splitlines():
            doc.add_paragraph(line)
    if sections.get("emails") is not None:
        doc.add_paragraph("Personal Emails:")
        for email in sections["emails"]:
            doc.add_paragraph(f"- {email}")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Basic extraction tests ────────────────────────────────────────────────────

def test_extract_single_talent_approved_response():
    from backend.services.docx_parser import extract_talent_sections

    docx = _make_docx(
        talent_name="Test Talent",
        approved="Hi there! Thanks for reaching out.",
    )
    result = extract_talent_sections(docx)
    assert "test talent" in result
    entry = result["test talent"]
    assert entry["full_name"] == "Test Talent"
    assert "Hi there" in entry["approved_response"]


def test_extract_personal_emails():
    from backend.services.docx_parser import extract_talent_sections

    docx = _make_docx(
        talent_name="Email Talent",
        approved="Thanks for the opportunity!",
        emails=["personal@gmail.com", "alt@yahoo.com"],
    )
    result = extract_talent_sections(docx)
    assert "email talent" in result
    assert set(result["email talent"]["personal_emails"]) == {"personal@gmail.com", "alt@yahoo.com"}


def test_extract_no_talents_returns_empty():
    from backend.services.docx_parser import extract_talent_sections

    doc = Document()
    doc.add_paragraph("This is some preamble text with no talent sections.")
    buf = BytesIO()
    doc.save(buf)

    result = extract_talent_sections(buf.getvalue())
    assert result == {}


def test_extract_multiple_talents():
    from backend.services.docx_parser import extract_talent_sections

    doc = Document()
    doc.add_paragraph("Talent: Alice")
    doc.add_paragraph("Approved Response:")
    doc.add_paragraph("Hello from Alice!")
    doc.add_paragraph("Talent: Bob")
    doc.add_paragraph("Approved Response:")
    doc.add_paragraph("Hello from Bob!")
    buf = BytesIO()
    doc.save(buf)

    result = extract_talent_sections(buf.getvalue())
    assert "alice" in result
    assert "bob" in result
    assert "Hello from Alice" in result["alice"]["approved_response"]
    assert "Hello from Bob" in result["bob"]["approved_response"]


def test_literal_markdown_is_preserved_verbatim():
    """The author types markdown as literal text — it must survive byte-for-byte.

    Regression guard. A previous parser re-emitted "**" for Word's own bold
    runs. Because the author's literal "**" was already in the text, that
    double-wrapped ("****fashion****") and the cleanup regexes then ate the
    opening delimiter, yielding "1 TikTok**" — an orphan marker that renders
    literally in the sent email. It corrupted 14 of 18 talents.

    The rate line below is the real authored shape, including the 5-space
    indent, and it is also formatted bold in Word — exactly the combination
    that broke before.
    """
    from backend.services.docx_parser import extract_talent_sections

    rate_line = "     **1 TikTok** [grayson.finks](https://www.tiktok.com/@grayson.finks) - $750"

    doc = Document()
    doc.add_paragraph("Talent: Bold Talent")
    doc.add_paragraph("Approved Response:")
    # A greeting always precedes the rate lines in the real doc; the block is
    # strip()ed as a whole, so the indent only survives on non-leading lines.
    doc.add_paragraph("Happy to share her rates below:")
    p = doc.add_paragraph()
    run = p.add_run(rate_line)
    run.bold = True  # Word bold ON TOP of the literal ** the author typed

    buf = BytesIO()
    doc.save(buf)

    ar = extract_talent_sections(buf.getvalue())["bold talent"]["approved_response"]

    assert ar.endswith(rate_line), f"expected literal text, got {ar!r}"
    assert "****" not in ar, "Word bold was re-emitted on top of literal markdown"
    assert ar.count("**") % 2 == 0, "unbalanced bold delimiters would render literally"


def test_leading_indent_on_rate_lines_is_preserved():
    """Leading whitespace is significant — rate lines carry a 5-space indent."""
    from backend.services.docx_parser import extract_talent_sections

    doc = Document()
    doc.add_paragraph("Talent: Indent Talent")
    doc.add_paragraph("Approved Response:")
    doc.add_paragraph("Rates below:")
    doc.add_paragraph("     **1 UGC Video** - $400")

    buf = BytesIO()
    doc.save(buf)

    ar = extract_talent_sections(buf.getvalue())["indent talent"]["approved_response"]
    assert "\n     **1 UGC Video** - $400" in ar, f"indent lost: {ar!r}"


def test_empty_docx_returns_empty():
    from backend.services.docx_parser import extract_talent_sections

    doc = Document()
    buf = BytesIO()
    doc.save(buf)

    result = extract_talent_sections(buf.getvalue())
    assert result == {}


# ── sop_versions provisioning ────────────────────────────────────────────────

def test_sop_fidelity_round_trip(_patch_sop_md_path):
    """SOP fidelity: parser → verbatim gate full round-trip.

    Builds a docx with a known approved response, parses it with
    extract_talent_sections, then asserts:
    - enforce_verbatim_response accepts the exact extracted text (returns None)
    - enforce_verbatim_response rejects a one-word modification (returns an error)

    This guards against parser regressions that would silently corrupt the
    approved response text and cause every auto-send to be blocked by the
    verbatim gate (or, worse, pass when they shouldn't).

    The autouse _patch_sop_md_path fixture is what makes KatrinaD's section
    available in the reply module's SOP cache. We need it injected here so we
    can call enforce_verbatim_response against a real SOP entry.
    """
    from backend.services.docx_parser import extract_talent_sections
    from backend.services.reply import enforce_verbatim_response

    # The approved response for KatrinaD in test_talents_sop.md fixture.
    # Deliberately uses the real fixture text so this test breaks if someone
    # silently changes the fixture's wording.
    known_approved = (
        "Thank you so much for reaching out about a potential partnership with Katrina D!\n\n"
        "Her rate is $150 per hour. Please let us know if you would like to move forward "
        "and we can discuss the scope!"
    )

    # Build a docx that mirrors the fixture structure
    doc = Document()
    doc.add_paragraph("Talent: Katrina D")
    doc.add_paragraph("Approved Response:")
    # Add each paragraph line that composes the approved response
    for line in known_approved.split("\n"):
        doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    parsed = extract_talent_sections(buf.getvalue())

    assert "katrina d" in parsed, "Parser must find 'Katrina D' section"
    extracted = parsed["katrina d"]["approved_response"]

    # The extracted text must match exactly what the verbatim gate will check against.
    # If it doesn't, every docx upload → auto-send would be dead on arrival.
    err = enforce_verbatim_response("Katrina D", extracted)
    assert err is None, (
        f"enforce_verbatim_response rejected its own parser output — "
        f"the round-trip is broken.\n"
        f"Extracted: {extracted!r}\n"
        f"Error: {err}"
    )

    # A one-word modification must be caught.
    modified = extracted.replace("Thank you so much", "Thanks so much")
    err_modified = enforce_verbatim_response("Katrina D", modified)
    assert err_modified is not None, (
        "enforce_verbatim_response should reject a modified draft but returned None"
    )


def test_ensure_sop_versions_table_creates_when_missing():
    """The SOP Manager must provision its own table.

    Production runs with SKIP_MIGRATIONS=true, so create_tables() never runs.
    sop_versions was consequently never created and every version-history call
    returned 500 UndefinedTable. _ensure_sop_versions_table() must create it on
    demand and be safe to call repeatedly.
    """
    from sqlalchemy import inspect
    from backend.models.db import SopVersion, get_engine
    from backend.routers.sop_admin import _ensure_sop_versions_table

    engine = get_engine()
    SopVersion.__table__.drop(bind=engine, checkfirst=True)
    assert "sop_versions" not in inspect(engine).get_table_names()

    _ensure_sop_versions_table()
    assert "sop_versions" in inspect(engine).get_table_names()

    cols = {c["name"] for c in inspect(engine).get_columns("sop_versions")}
    assert "doc_type" in cols

    # Idempotent — a second call must not raise.
    _ensure_sop_versions_table()


# ── Personal-email bullet formats ─────────────────────────────────────────────

@pytest.mark.parametrize("prefix,label", [("- ", "hyphen bullet"), ("• ", "unicode bullet"), ("", "bare line")])
def test_personal_emails_parse_with_or_without_bullets(prefix, label):
    """SOP revisions disagree on the bullet: v15 used '- addr', v15-c uses 'addr'.

    Requiring the bullet made v15-c parse to zero emails for all 18 talents,
    which would silently freeze Scenario C personal-email routing.
    """
    from backend.services.docx_parser import extract_talent_sections

    doc = Document()
    doc.add_paragraph("Talent: Bullet Talent")
    doc.add_paragraph("Approved Response:")
    doc.add_paragraph("Hello.")
    doc.add_paragraph("Scenario C: Personal Email Forward")
    doc.add_paragraph("Personal Emails: ")
    doc.add_paragraph(f"{prefix}one@example.com")
    doc.add_paragraph(f"{prefix}two@example.com")
    buf = BytesIO()
    doc.save(buf)

    got = extract_talent_sections(buf.getvalue())["bullet talent"]["personal_emails"]
    assert got == ["one@example.com", "two@example.com"], f"{label} form failed: {got}"


def test_non_email_line_still_ends_the_email_list():
    """Loosening the bullet must not let prose leak into personal_emails."""
    from backend.services.docx_parser import extract_talent_sections

    doc = Document()
    doc.add_paragraph("Talent: Boundary Talent")
    doc.add_paragraph("Approved Response:")
    doc.add_paragraph("Hello.")
    doc.add_paragraph("Scenario C: Personal Email Forward")
    doc.add_paragraph("Personal Emails:")
    doc.add_paragraph("only@example.com")
    doc.add_paragraph("Some trailing prose that is not an email.")
    doc.add_paragraph("later@example.com")
    buf = BytesIO()
    doc.save(buf)

    got = extract_talent_sections(buf.getvalue())["boundary talent"]["personal_emails"]
    assert got == ["only@example.com"], f"prose or post-prose lines leaked: {got}"
