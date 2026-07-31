"""SOP Admin router — manage talent data without directly editing sop.md."""
from __future__ import annotations

import logging
import re
import urllib.request
from datetime import datetime
from io import BytesIO
from pathlib import Path

from fastapi import APIRouter, Depends, File, Form, HTTPException, Response, UploadFile
from pydantic import BaseModel

from backend.core.config import get_settings
from backend.routers.deps import verify_api_key
from backend.services import sop_writer as _writer
from backend.services.sop_parser import parse_sop_md, validate_profiles

logger = logging.getLogger(__name__)

_SOP_PATH = Path(__file__).resolve().parents[2] / "sheets" / "sop.md"

# Scenario headings other than A. extract_talent_sections() only pulls Scenario A,
# so any of these present in an uploaded docx will NOT reach sop.md.
_NON_A_SCENARIO_RE = re.compile(r"^[ \t]*Scenario\s+([B-Z])\b[^\n]*", re.IGNORECASE)

_TALENT_HEADING_RE = re.compile(
    r"^[ \t]*(?:#+[ \t]*)?Talent:[ \t]*(?P<name>[^\r\n]*)[ \t]*$",
    re.MULTILINE,
)

router = APIRouter(
    prefix="/admin",
    tags=["admin"],
    include_in_schema=False,
)


def _read_sop() -> str:
    return _SOP_PATH.read_text(encoding="utf-8")


def _detect_unmerged_scenarios(docx_bytes: bytes, docx_talents: dict) -> list[str]:
    """Warn about Scenario B/D/… blocks in the docx that the merge will not apply.

    extract_talent_sections() returns Scenario A only. Production sop.md happens
    to contain just Scenario A and C today, so nothing is lost — but a Scenario B
    (the rate-negotiation counter-offer) added to a future SOP revision would be
    dropped with no indication anywhere. Surface it instead.

    Returns one message per (talent, scenario letter). Never raises: a failure to
    warn must not fail the upload.
    """
    try:
        from docx import Document
    except ImportError:
        return []

    try:
        paragraphs = [p.text for p in Document(BytesIO(docx_bytes)).paragraphs]
    except Exception:
        return []

    messages: list[str] = []
    current = None
    seen: set[tuple[str, str]] = set()
    for raw in paragraphs:
        talent_match = _TALENT_HEADING_RE.match(raw)
        if talent_match:
            current = talent_match.group("name").strip()
            continue
        scenario_match = _NON_A_SCENARIO_RE.match(raw)
        if scenario_match and current:
            letter = scenario_match.group(1).upper()
            # Scenario C is personal-email routing, which IS handled separately.
            if letter == "C":
                continue
            if (current, letter) in seen:
                continue
            seen.add((current, letter))
            messages.append(
                f"Scenario {letter} for {current} was found in the docx but NOT merged "
                f"— only Scenario A is applied. Update sop.md by hand if this is needed."
            )
    return messages


def _ensure_sop_versions_table() -> None:
    """Create sop_versions if it is missing, then add any new columns.

    This service runs with SKIP_MIGRATIONS=true, so create_tables() never runs
    and no schema change applies itself on deploy. sop_versions was therefore
    never created at all, and every version-history call 500'd with
    UndefinedTable. Rather than depend on the migration path, the SOP Manager
    provisions its own table on demand.

    Both operations are additive and idempotent: create() uses checkfirst, and
    the ADD COLUMN is IF NOT EXISTS. Nothing here can drop or rewrite data.
    """
    from sqlalchemy import inspect, text
    from backend.models.db import SopVersion, get_engine

    engine = get_engine()
    SopVersion.__table__.create(bind=engine, checkfirst=True)

    # An older deploy may have created the table before doc_type existed.
    try:
        cols = {c["name"] for c in inspect(engine).get_columns("sop_versions")}
    except Exception:
        return
    if "doc_type" not in cols and engine.dialect.name != "sqlite":
        with engine.connect() as conn:
            conn.execute(text(
                "ALTER TABLE sop_versions "
                "ADD COLUMN IF NOT EXISTS doc_type VARCHAR(32) NOT NULL DEFAULT 'sop'"
            ))
            conn.commit()


def _extract_approved_response(sop_text: str, talent_key: str) -> str:
    """Extract the approved response text for a talent from sop.md text."""
    matches = list(_TALENT_HEADING_RE.finditer(sop_text))
    for i, match in enumerate(matches):
        section_start = match.start()
        section_end = matches[i + 1].start() if i + 1 < len(matches) else len(sop_text)
        section = sop_text[section_start:section_end]
        key_m = re.search(r"^[ \t]*Key[ \t]*:[ \t]*(.+)$", section, re.MULTILINE)
        if not key_m or key_m.group(1).strip().lower() != talent_key.lower():
            continue
        ar_m = re.search(r"^[ \t]*Approved Response:[ \t]*$", section, re.MULTILINE)
        if not ar_m:
            return ""
        ar_end = ar_m.end()
        next_scenario = re.search(r"^[ \t]*Scenario\b", section[ar_end:], re.MULTILINE)
        content_end = ar_end + next_scenario.start() if next_scenario else len(section)
        return section[ar_end:content_end].strip()
    return ""


def _resolve_profile(talent_key: str):
    """Return the TalentProfile for the given key (case-insensitive). 404 if not found."""
    profiles = get_settings().talent_profiles
    profile = profiles.get(talent_key) or next(
        (p for k, p in profiles.items() if k.lower() == talent_key.lower()), None
    )
    if profile is None:
        raise HTTPException(status_code=404, detail=f"Talent '{talent_key}' not found")
    return profile


def _persist_active_sop_version(sop_text: str, label: str) -> int | None:
    """Save sop_text as the new active 'sop' SopVersion row.

    Every Render deploy starts from a fresh filesystem and restores sheets/sop.md
    from whichever SopVersion row is currently active (see main.py on_startup).
    Any live edit that only calls sop_writer.write_sop_md() — e.g. update_talent,
    toggle_auto_send — changes disk but NOT the DB, so it silently reverts on the
    next deploy. This mirrors the persistence upload_sop_v2 already does, so every
    write path keeps disk and DB in sync. Non-fatal: the live edit already
    succeeded on disk even if this fails.
    """
    try:
        from backend.models.db import SopVersion, get_session_factory
        _ensure_sop_versions_table()
        db = get_session_factory()()
        try:
            db.query(SopVersion).filter(
                SopVersion.is_active == True,  # noqa: E712
                SopVersion.doc_type == "sop",
            ).update({"is_active": False})
            new_ver = SopVersion(
                version_label=label,
                raw_content=sop_text,
                talent_count=len(parse_sop_md(sop_text)),
                is_active=True,
                doc_type="sop",
            )
            db.add(new_ver)
            db.commit()
            db.refresh(new_ver)
            return new_ver.id
        finally:
            db.close()
    except Exception:
        logger.warning("Live SOP edit: DB version persist failed (non-fatal — sop.md is still correct on disk)", exc_info=True)
        return None


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.get("/api/talents", dependencies=[Depends(verify_api_key)])
def list_talents():
    profiles = get_settings().talent_profiles
    return {
        "talents": [
            {
                "key": p.key,
                "full_name": p.full_name,
                "manager": p.manager,
                "manager_email": p.manager_email,
                "minimum_rate_usd": p.minimum_rate_usd,
                "rate_unit": p.rate_unit,
                "auto_send": p.auto_send,
                "paused": p.paused,
                "has_approved_response": p.has_approved_response,
                "personal_emails": p.personal_emails,
            }
            for p in profiles.values()
        ]
    }


@router.get("/api/talents/{talent_key}", dependencies=[Depends(verify_api_key)])
def get_talent(talent_key: str):
    profile = _resolve_profile(talent_key)
    sop_text = _read_sop()
    approved_response = _extract_approved_response(sop_text, profile.key)
    return {
        "key": profile.key,
        "full_name": profile.full_name,
        "manager": profile.manager,
        "manager_email": profile.manager_email,
        "gmail_connection_name": profile.gmail_connection_name,
        "minimum_rate_usd": profile.minimum_rate_usd,
        "rate_unit": profile.rate_unit,
        "auto_send": profile.auto_send,
        "paused": profile.paused,
        "has_approved_response": profile.has_approved_response,
        "personal_emails": profile.personal_emails,
        "approved_response": approved_response,
    }


class TalentUpdateRequest(BaseModel):
    minimum_rate_usd: int | None = None
    rate_unit: str | None = None
    auto_send: bool | None = None
    paused: bool | None = None
    approved_response: str | None = None
    personal_emails: list[str] | None = None
    manager: str | None = None


@router.put("/api/talents/{talent_key}", dependencies=[Depends(verify_api_key)])
def update_talent(talent_key: str, body: TalentUpdateRequest):
    profile = _resolve_profile(talent_key)

    errors = _writer.validate_before_write(
        body.minimum_rate_usd, body.personal_emails, body.approved_response
    )
    if errors:
        raise HTTPException(status_code=400, detail={"errors": errors})

    sop_text = _read_sop()

    if body.minimum_rate_usd is not None or body.rate_unit is not None:
        rate_usd = body.minimum_rate_usd if body.minimum_rate_usd is not None else profile.minimum_rate_usd
        rate_unit = body.rate_unit if body.rate_unit is not None else profile.rate_unit
        sop_text = _writer.update_talent_field(
            sop_text, profile.key, "Min Rate",
            f"${rate_usd} {rate_unit}".strip()
        )

    if body.auto_send is not None:
        sop_text = _writer.update_talent_field(
            sop_text, profile.key, "Auto Send", "yes" if body.auto_send else "no"
        )

    if body.paused is not None:
        sop_text = _writer.update_talent_field(
            sop_text, profile.key, "Paused", "yes" if body.paused else "no"
        )

    if body.manager is not None:
        sop_text = _writer.update_talent_field(sop_text, profile.key, "Manager", body.manager)

    if body.approved_response is not None:
        sop_text = _writer.update_approved_response(sop_text, profile.key, body.approved_response)

    if body.personal_emails is not None:
        sop_text = _writer.update_personal_emails(sop_text, profile.key, body.personal_emails)

    _writer.write_sop_md(sop_text)
    version_id = _persist_active_sop_version(
        sop_text, f"Live edit — {profile.key} — {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
    )
    return {"status": "ok", "key": profile.key, "version_id": version_id}


@router.post("/api/talents/{talent_key}/toggle-auto-send", dependencies=[Depends(verify_api_key)])
def toggle_auto_send(talent_key: str):
    profile = _resolve_profile(talent_key)
    new_value = not profile.auto_send
    sop_text = _read_sop()
    sop_text = _writer.update_talent_field(
        sop_text, profile.key, "Auto Send", "yes" if new_value else "no"
    )
    _writer.write_sop_md(sop_text)
    version_id = _persist_active_sop_version(
        sop_text, f"Live edit — {profile.key} auto-send toggle — {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
    )
    return {"key": profile.key, "auto_send": new_value, "version_id": version_id}


@router.get("/api/sop/raw", dependencies=[Depends(verify_api_key)])
def sop_raw():
    return Response(content=_read_sop(), media_type="text/plain")


@router.post("/api/sop/promote-repo-version", dependencies=[Depends(verify_api_key)])
def promote_repo_version():
    """
    Make the CURRENT on-disk sheets/sop.md (i.e. whatever git deployed) the new
    active DB version, so it survives the next redeploy instead of being silently
    overwritten by an older active SopVersion row from a prior docx upload.

    Use this once after any git-committed sop.md change, right after it deploys —
    otherwise main.py's startup restore will keep reverting to the stale DB row
    on every subsequent boot, even though the repo file is correct.
    """
    sop_text = _read_sop()
    profiles = parse_sop_md(sop_text)
    version_id = _persist_active_sop_version(
        sop_text, f"Promoted from repo — {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
    )
    return {
        "status": "ok" if version_id else "disk_only",
        "version_id": version_id,
        "talent_count": len(profiles),
        "talent_keys": sorted(profiles.keys()),
    }


@router.post("/api/sop/import-docx", dependencies=[Depends(verify_api_key)])
async def import_sop_docx(file: UploadFile = File(...)):
    """Parse a .docx file and return a preview — does not write anything."""
    try:
        from docx import Document  # python-docx
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed on server")

    content = await file.read()
    try:
        doc = Document(BytesIO(content))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not parse docx: {exc}")

    sop_text = "\n".join(p.text for p in doc.paragraphs)
    profiles = parse_sop_md(sop_text)

    if len(profiles) < 5:
        raise HTTPException(
            status_code=400,
            detail=f"Only {len(profiles)} talent profile(s) parsed — check that the docx uses the same format as sop.md (Talent:, Key:, Manager:, etc.)",
        )

    warnings = validate_profiles(profiles)
    return {
        "talent_count": len(profiles),
        "talent_names": [p.full_name for p in profiles.values()],
        "warnings": warnings,
        "sop_text": sop_text,
    }


# ── V2 SOP upload (rich-text aware) ──────────────────────────────────────────


@router.post("/api/sop/upload-v2", dependencies=[Depends(verify_api_key)])
async def upload_sop_v2(
    file: UploadFile = File(...),
    label: str = Form(""),
):
    """
    Accept a .docx SOP file, extract rich text (preserving **bold** and
    [hyperlinks](url)), merge the approved responses + personal emails into
    the existing sop.md (keeping all metadata: Key, Gmail, Min Rate, etc.),
    write the result to disk, persist it in the sop_versions table, and
    clear all in-memory caches so the new SOP is live immediately.

    Does NOT trigger a Render redeploy — that's a separate button.
    The startup handler restores the active DB version on the next deploy.
    """
    # NB: guarding the module import alone is not enough — docx_parser imports
    # python-docx *inside* extract_talent_sections(), so this import succeeds
    # even when the dependency is missing and the ImportError surfaces at call
    # time. It then got swallowed by the broad handler below and reported as
    # 400 "Could not parse docx", blaming the operator's file for a missing
    # server dependency. Catch ImportError separately at the call site.
    try:
        from backend.services.docx_parser import extract_talent_sections
    except ImportError as exc:
        raise HTTPException(status_code=500, detail=f"docx_parser not available: {exc}")

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    try:
        docx_talents = extract_talent_sections(content)
    except ImportError as exc:
        logger.error("SOP upload failed — python-docx not installed on the server: %s", exc)
        raise HTTPException(
            status_code=500,
            detail=(
                "Server is missing the python-docx dependency — this is not a problem "
                f"with your file. Redeploy so requirements.txt installs. ({exc})"
            ),
        )
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not parse docx: {exc}")

    if not docx_talents:
        raise HTTPException(
            status_code=400,
            detail="No 'Talent: Name' sections found in the docx — check the file format",
        )

    # Read current sop.md and parse existing profiles
    sop_text = _read_sop()
    existing_profiles = parse_sop_md(sop_text)

    matched: list[str] = []
    unmatched_sop: list[str] = []
    unmatched_docx: list[str] = []
    warnings: list[str] = []

    for key, profile in existing_profiles.items():
        name_lower = profile.full_name.lower()
        docx_entry = docx_talents.get(name_lower)

        # Fallback: match by first word of name (catches "Brittany" matching "Brittany Kuhl")
        if docx_entry is None:
            first_word = name_lower.split()[0] if name_lower.split() else ""
            docx_entry = next(
                (v for k, v in docx_talents.items() if k.startswith(first_word + " ") or k == first_word),
                None,
            )

        if docx_entry is None:
            unmatched_sop.append(profile.key)
            continue

        matched.append(profile.key)

        if docx_entry.get("approved_response"):
            try:
                sop_text = _writer.update_approved_response(
                    sop_text, profile.key, docx_entry["approved_response"]
                )
            except ValueError as exc:
                warnings.append(f"{profile.key}: approved response not updated — {exc}")

        if docx_entry.get("personal_emails"):
            try:
                sop_text = _writer.update_personal_emails(
                    sop_text, profile.key, docx_entry["personal_emails"]
                )
            except ValueError as exc:
                warnings.append(f"{profile.key}: personal emails not updated — {exc}")

    # Find docx talents that had no match in sop.md. These are skipped on
    # purpose: adding a talent needs metadata the docx does not carry (Key,
    # Gmail, Min Rate, Auto Send, Paused), so a silent insert would create a
    # half-configured profile. Log it so it shows up in Render logs, not just
    # in the browser response.
    matched_name_lowers = {existing_profiles[k].full_name.lower() for k in matched}
    for docx_name, docx_entry in docx_talents.items():
        if docx_name not in matched_name_lowers:
            unmatched_docx.append(docx_entry["full_name"])
            logger.warning(
                "SOP upload: '%s' is in the docx but not in sop.md — skipped. "
                "New talents need manual onboarding (Key/Gmail/Min Rate/Auto Send/Paused).",
                docx_entry["full_name"],
            )

    # Scenario A is the only block the parser extracts. Anything else in the
    # docx is silently ignored by the merge, so say so loudly rather than
    # letting a counter-offer response quietly fail to ship.
    for scenario_warning in _detect_unmerged_scenarios(content, docx_talents):
        warnings.append(scenario_warning)
        logger.warning("SOP upload: %s", scenario_warning)

    # Write updated sop.md to disk + clear caches
    _writer.write_sop_md(sop_text)

    # Persist merged sop.md to sop_versions table as new active version
    version_id: int | None = None
    try:
        from backend.models.db import SopVersion, get_session_factory
        _ensure_sop_versions_table()
        db = get_session_factory()()
        try:
            # Scope to doc_type — the workflow doc keeps its own active version.
            db.query(SopVersion).filter(
                SopVersion.is_active == True,  # noqa: E712
                SopVersion.doc_type == "sop",
            ).update({"is_active": False})
            version_label = label.strip() or f"Upload {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
            new_ver = SopVersion(
                version_label=version_label,
                raw_content=sop_text,
                talent_count=len(existing_profiles),
                is_active=True,
                doc_type="sop",
            )
            db.add(new_ver)
            db.commit()
            db.refresh(new_ver)
            version_id = new_ver.id
        finally:
            db.close()
    except Exception as exc:
        warnings.append(f"DB version save failed (non-fatal — SOP is still live on disk): {exc}")

    return {
        "status": "ok",
        "matched": matched,
        "unmatched_sop": unmatched_sop,   # sop.md talents with no docx match (response unchanged)
        "unmatched_docx": unmatched_docx,  # docx talents with no sop.md match (ignored)
        "warnings": warnings,
        "version_id": version_id,
        "talent_count": len(existing_profiles),
    }


@router.post("/api/workflow/upload", dependencies=[Depends(verify_api_key)])
async def upload_workflow(
    file: UploadFile = File(...),
    label: str = Form(""),
    confirm: bool = Form(False),
):
    """
    Accept a .docx of the Automated Send Workflow.

    Two-step by design: the first call returns a unified diff and writes
    nothing. Re-submit the same file with confirm=true to apply it.

    Unlike sop.md this document carries no machine-read metadata, so a whole
    document replace is safe — but it is still validated for the expected
    section anchors, so an empty file, or the SOP dropped into the wrong box,
    is rejected before anything is written.
    """
    import difflib

    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(status_code=500, detail=f"python-docx not available: {exc}")

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    try:
        doc = Document(BytesIO(content))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not parse docx: {exc}")

    # paragraph.text is correct here for the same reason it is for the SOP:
    # the document's markup is typed as literal text, not Word formatting.
    new_text = "\n".join(p.text for p in doc.paragraphs).strip() + "\n"

    problems = _writer.validate_workflow_text(new_text)
    if problems:
        raise HTTPException(status_code=400, detail="; ".join(problems))

    current = _writer.read_workflow_md()
    diff = list(
        difflib.unified_diff(
            current.splitlines(), new_text.splitlines(),
            fromfile="current", tofile="uploaded", lineterm="", n=2,
        )
    )
    added = sum(1 for d in diff if d.startswith("+") and not d.startswith("+++"))
    removed = sum(1 for d in diff if d.startswith("-") and not d.startswith("---"))

    if not confirm:
        return {
            "status": "preview",
            "changed": bool(diff),
            "lines_added": added,
            "lines_removed": removed,
            "diff": diff[:400],
            "message": "Nothing written. Re-submit with confirm to apply.",
        }

    _writer.write_workflow_md(new_text)

    version_id: int | None = None
    save_warning: str | None = None
    try:
        from backend.models.db import SopVersion, get_session_factory
        _ensure_sop_versions_table()
        db = get_session_factory()()
        try:
            db.query(SopVersion).filter(
                SopVersion.is_active == True,  # noqa: E712
                SopVersion.doc_type == "workflow",
            ).update({"is_active": False})
            new_ver = SopVersion(
                version_label=label.strip() or f"Workflow {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
                raw_content=new_text,
                talent_count=None,
                is_active=True,
                doc_type="workflow",
            )
            db.add(new_ver)
            db.commit()
            db.refresh(new_ver)
            version_id = new_ver.id
        finally:
            db.close()
    except Exception as exc:
        save_warning = f"DB version save failed (non-fatal — workflow is live on disk): {exc}"

    return {
        "status": "ok",
        "lines_added": added,
        "lines_removed": removed,
        "version_id": version_id,
        "warnings": [save_warning] if save_warning else [],
    }


# ── Version history ───────────────────────────────────────────────────────────


@router.get("/api/sop/versions", dependencies=[Depends(verify_api_key)])
def list_sop_versions():
    """Return the 20 most recent SOP versions from the DB."""
    try:
        from backend.models.db import SopVersion, get_session_factory
    except ImportError as exc:
        raise HTTPException(status_code=500, detail=str(exc))

    try:
        _ensure_sop_versions_table()
    except Exception as exc:  # provisioning failure must not 500 the page
        raise HTTPException(status_code=503, detail=f"Version store unavailable: {exc}")

    db = get_session_factory()()
    try:
        versions = (
            db.query(SopVersion)
            .order_by(SopVersion.uploaded_at.desc())
            .limit(20)
            .all()
        )
        return {
            "versions": [
                {
                    "id": v.id,
                    "label": v.version_label,
                    "talent_count": v.talent_count,
                    "uploaded_at": v.uploaded_at.isoformat() if v.uploaded_at else None,
                    "is_active": v.is_active,
                    "doc_type": getattr(v, "doc_type", "sop") or "sop",
                }
                for v in versions
            ]
        }
    finally:
        db.close()


@router.post("/api/sop/versions/{version_id}/restore", dependencies=[Depends(verify_api_key)])
def restore_sop_version(version_id: int):
    """Restore a past SOP version: write it to sop.md, mark it active in DB."""
    try:
        from backend.models.db import SopVersion, get_session_factory
    except ImportError as exc:
        raise HTTPException(status_code=500, detail=str(exc))

    try:
        _ensure_sop_versions_table()
    except Exception as exc:  # provisioning failure must not 500 the page
        raise HTTPException(status_code=503, detail=f"Version store unavailable: {exc}")

    db = get_session_factory()()
    try:
        version = db.query(SopVersion).filter(SopVersion.id == version_id).first()
        if version is None:
            raise HTTPException(status_code=404, detail=f"Version {version_id} not found")

        # Route by document type — restoring a workflow version must not be
        # written into sop.md, and vice versa.
        doc_type = getattr(version, "doc_type", "sop") or "sop"
        if doc_type == "workflow":
            _writer.write_workflow_md(version.raw_content)
        else:
            _writer.write_sop_md(version.raw_content)

        # Mark active within this document type only
        db.query(SopVersion).filter(
            SopVersion.is_active == True,  # noqa: E712
            SopVersion.doc_type == doc_type,
        ).update({"is_active": False})
        version.is_active = True
        db.commit()

        return {
            "status": "ok",
            "version_id": version_id,
            "label": version.version_label,
            "doc_type": doc_type,
        }
    finally:
        db.close()


# ── Render deploy hook ────────────────────────────────────────────────────────


@router.post("/api/sop/deploy-render", dependencies=[Depends(verify_api_key)])
def deploy_to_render():
    """
    Trigger a Render redeploy via the deploy hook URL stored in
    RENDER_DEPLOY_HOOK_URL env var.  The SOP itself is already live on disk
    (write_sop_md is instant); this deploy is only needed so the next startup
    also restores from DB correctly after any code changes.
    """
    hook = get_settings().render_deploy_hook_url
    if not hook:
        raise HTTPException(
            status_code=400,
            detail="RENDER_DEPLOY_HOOK_URL not set — add it in Render → Environment Variables",
        )
    try:
        req = urllib.request.Request(hook, method="POST")
        with urllib.request.urlopen(req, timeout=15) as resp:
            return {"status": "ok", "http_status": resp.status}
    except urllib.error.HTTPError as exc:
        raise HTTPException(status_code=502, detail=f"Render hook returned {exc.code}: {exc.reason}")
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Render deploy hook failed: {exc}")
