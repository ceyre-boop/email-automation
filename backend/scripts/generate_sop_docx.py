"""
Generate sheets/sop.docx from sheets/sop_data.json.

This script builds the authoritative SOP document used as:
  1. AI agent context (backend reads the .docx via /api/dashboard/sop-html)
  2. Human-readable reference (opens in Word)

Run from the repo root:
    python -m backend.scripts.generate_sop_docx
"""
from __future__ import annotations

import json
import pathlib

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Paths ─────────────────────────────────────────────────────────────────────
REPO_ROOT = pathlib.Path(__file__).resolve().parents[2]
SOP_JSON = REPO_ROOT / "sheets" / "sop_data.json"
OUT_DOCX = REPO_ROOT / "sheets" / "sop.docx"

# ── Action-rule keywords (same filter as reply.py._build_sop_rules_text) ─────
_ACTION_KEYWORDS = (
    "move to", "cc ", "delete", "ask for consult",
    "tagged email", "marked a initial",
)

def _is_action_rule(response: str) -> bool:
    r = response.lower().strip()
    return (
        any(r.startswith(kw) for kw in _ACTION_KEYWORDS)
        or any(kw in r for kw in ("move to ", "cc cara", "cc chenni", "cc nicole", "move it to"))
    )

def _derive_internal_note(rules: list[dict], idx: int) -> str:
    """
    Look at the rule immediately after idx. If it's an action rule whose
    trigger overlaps the current context, return it as an internal note.
    """
    if idx + 1 < len(rules):
        nxt = rules[idx + 1]
        resp = nxt.get("response", "").strip()
        if _is_action_rule(resp):
            trigger = nxt.get("trigger", "").replace("\n", " ").strip()
            return f"{trigger} → {resp}"
    return "—"

def _add_bold_label(doc: Document, label: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True

def _add_response_paragraphs(doc: Document, text: str) -> None:
    """Split response on newlines and add each line as its own paragraph."""
    for line in text.split("\n"):
        doc.add_paragraph(line.strip() if line.strip() else "")

# ── Global rules section ──────────────────────────────────────────────────────
GLOBAL_RULES = [
    "The SOP document must be followed explicitly. Do not deviate from approved responses. "
    "Do not rewrite, improve, shorten, expand, or personalize approved responses unless "
    "specifically instructed by an admin.",

    "Talent matching is mandatory. Each talent has different rates, terms, and response "
    "language. Always identify the correct talent before selecting a response. Never use "
    "one talent's response for another talent.",

    "This workflow is for INITIAL inbound emails only. Draft responses only for first-time "
    "inbound emails or new deal inquiries. If the email is part of an ongoing thread, "
    "follow-up, negotiation, or reply after the initial response, do not draft a response. "
    "Instead return: Classification: Human Admin Required — Reason: This appears to be a "
    "follow-up or ongoing conversation.",

    "Default to the Initial Approved Response. Each talent has an Initial Approved Response "
    "which is the default for valid inbound opportunities. Only choose another approved "
    "response if the email clearly matches a more specific scenario.",

    "Err on the side of responding. Only classify as Spam/Trash/Archive if the email is "
    "clearly and truly spam. If there is any reasonable chance the email is a real brand, "
    "agency, PR, event, collaboration, gifting, partnership, or paid inquiry, use the "
    "talent's Initial Approved Response. It is better to reply to a questionable email "
    "than to accidentally ignore a real opportunity.",

    "Spam handling must be conservative. Do not classify as Spam merely because the email "
    "is vague, low-budget, generic, poorly written, or from an unfamiliar sender. "
    "Non-English emails (e.g. Chinese market emails) are NOT automatically spam — treat "
    "them as Score 2 if they reference a brand or collaboration context. Classify as Spam "
    "only when there are clear indicators: phishing, scams, suspicious links, unrelated "
    "services, mass SEO/web/design pitches, fake invoices, malware, adult/illegal content.",

    "Output must clearly state the action taken: Approved Response / Human Admin Required "
    "/ Spam / Ignore.",

    "If using an approved response: return the exact approved response only. Do not modify "
    "the response text. Do not combine multiple approved responses. Do not add extra "
    "commentary inside the email draft.",
]

KEY_TERMS = {
    "Initial Approved Response": (
        "The default email reply template for a talent. Used when a brand reaches out for "
        "the first time with no specific offer, or when asking about rates/collabs."
    ),
    "Human Admin Required": (
        "Classification for emails that are part of an ongoing conversation, follow-up, "
        "or negotiation. The AI does NOT draft a reply — a human manager handles it."
    ),
    "Spam": (
        "Emails that are clearly junk: phishing, scams, unrelated service pitches, fake "
        "prizes, malware. Conservative threshold — when in doubt, do not classify as Spam."
    ),
    "Ignore": (
        "Emails that are not real brand deals, are fan mail, or require no response. "
        "Trinity receives fan mail which is classified as Ignore."
    ),
    "Score 1 / Archive": "Email is definitively trash — archived immediately, no reply.",
    "Score 2 / Flag": "Email is uncertain or a follow-up — flagged for human review, no draft.",
    "Score 3 / Draft": "Email is a real opportunity — AI drafts a reply using the approved SOP response.",
}

# ── Main ──────────────────────────────────────────────────────────────────────

def main() -> None:
    sop_data: dict = json.loads(SOP_JSON.read_text(encoding="utf-8"))
    doc = Document()

    # ── Title ──
    title = doc.add_heading("TABOOST — Talent Inbox SOP", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── SOP Status Index ──
    doc.add_heading("SOP STATUS — WHO IS READY", level=1)
    approved = [k for k, v in sop_data.items() if v.get("sop_status") == "approved"]
    pending  = [k for k, v in sop_data.items() if v.get("sop_status") != "approved"]
    p = doc.add_paragraph()
    p.add_run("✅ APPROVED (AI may draft): ").bold = True
    p.add_run(", ".join(v.get("full_name", k) for k, v in sop_data.items() if k in approved))
    p2 = doc.add_paragraph()
    p2.add_run("⏳ PENDING (AI will NOT draft — Human Admin Required): ").bold = True
    p2.add_run(", ".join(v.get("full_name", k) for k, v in sop_data.items() if k in pending))

    # ── Global Rules ──
    doc.add_heading("GLOBAL RULES — MANDATORY", level=1)
    for i, rule in enumerate(GLOBAL_RULES, 1):
        doc.add_paragraph(f"{i}. {rule}")

    # ── Key Terms ──
    doc.add_heading("KEY TERMS", level=1)
    for term, definition in KEY_TERMS.items():
        p = doc.add_paragraph()
        run = p.add_run(f"{term}: ")
        run.bold = True
        p.add_run(definition)

    # ── Per-talent sections ──
    for talent_key, talent_data in sop_data.items():
        full_name = talent_data.get("full_name", talent_key)
        manager = talent_data.get("manager", "")
        manager_email = talent_data.get("manager_email", "")
        min_rate = talent_data.get("min_rate_usd", 0)
        rate_unit = talent_data.get("rate_unit", "per video")
        rules = talent_data.get("rules", [])
        status = talent_data.get("sop_status", "pending")

        doc.add_heading(full_name.upper(), level=1)

        # Status banner
        status_p = doc.add_paragraph()
        if status == "approved":
            status_p.add_run("✅ SOP APPROVED — AI may draft for this talent").bold = True
        else:
            status_p.add_run("⏳ SOP PENDING — AI will NOT draft. All emails routed to Human Admin Required.").bold = True

        # Talent metadata
        meta = doc.add_paragraph()
        mgr_str = f"{manager} ({manager_email})" if manager_email else manager
        meta.add_run(f"Manager: {mgr_str}  |  Minimum Rate: ${min_rate} {rate_unit}")

        if status != "approved":
            doc.add_paragraph("(Add approved scenarios below when ready.)")
            continue

        # ── New scenario format ──
        if rules and "scenario" in rules[0]:
            for rule in rules:
                scenario = rule.get("scenario", "")
                label = rule.get("label", "")
                is_default = rule.get("is_default", False)
                default_tag = " — DEFAULT RESPONSE" if is_default else ""
                doc.add_heading(f"Scenario {scenario}: {label}{default_tag}", level=2)

                _add_bold_label(doc, "Use when:")
                for uw in rule.get("use_when", []):
                    doc.add_paragraph(f"• {uw}")

                _add_bold_label(doc, "Do not use when:")
                for dnuw in rule.get("do_not_use_when", []):
                    doc.add_paragraph(f"• {dnuw}")

                cc = rule.get("cc")
                if cc:
                    _add_bold_label(doc, f"CC: {cc}")

                _add_bold_label(doc, "Approved Response:")
                _add_response_paragraphs(doc, rule.get("response", ""))
            continue

        # ── Legacy format fallback ──
        scenario_num = 0
        for idx, rule in enumerate(rules):
            trigger = rule.get("trigger", "").replace("\r\n", "\n").strip()
            response = rule.get("response", "").replace("\r\n", "\n").strip()

            if _is_action_rule(response):
                continue
            if len(response) < 20 and "manager:" in response.lower():
                continue

            scenario_num += 1
            trigger_short = trigger.split("\n")[0][:60]
            doc.add_heading(f"SCENARIO {scenario_num}: {trigger_short}", level=2)

            _add_bold_label(doc, "Use when (triggers):")
            doc.add_paragraph(trigger)

            _add_bold_label(doc, "Do not use when (exceptions):")
            doc.add_paragraph("—")

            _add_bold_label(doc, "Approved Response:")
            _add_response_paragraphs(doc, response)

            _add_bold_label(doc, "Internal Note:")
            doc.add_paragraph(_derive_internal_note(rules, idx))

        if scenario_num == 0:
            doc.add_paragraph("(No approved responses yet — add them here.)")

    doc.save(OUT_DOCX)
    print(f"✓ SOP document written to {OUT_DOCX}")
    print(f"  Talents: {len(sop_data)}")


if __name__ == "__main__":
    main()
